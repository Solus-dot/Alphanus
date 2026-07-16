from __future__ import annotations

import csv
from pathlib import Path
from typing import Any

from skills.runtime import ToolExecutionEnv

_DEFAULT_MAX_CHARS = 20_000
_MAX_CHARS_LIMIT = 1_000_000
_DEFAULT_MAX_ROWS = 200
_MAX_ROWS_LIMIT = 5000

TOOL_SPEC_ROWS = {  # fmt: skip
    "extract_document_text": ("document_read", False, ("read",), "Extract text from TXT, CSV, PDF, or DOCX. PDF/DOCX require optional dependencies.", {"path": {"type": "string"}, "max_chars": {"type": "integer"}}, ("path",), False),
    "extract_document_tables": ("document_read", False, ("read", "list"), "Extract table-like rows from CSV or DOCX tables when optional dependencies are available.", {"path": {"type": "string"}, "max_rows": {"type": "integer"}}, ("path",), False),
}
def _ok(data: dict[str, object]) -> dict[str, object]:
    return {"ok": True, "data": data, "error": None, "meta": {}}


def _err(code: str, message: str, data: dict[str, object] | None = None) -> dict[str, object]:
    return {"ok": False, "data": data or {}, "error": {"code": code, "message": message}, "meta": {}}


def _resolve_path(raw: str, env: ToolExecutionEnv) -> Path:
    if not raw.strip():
        raise ValueError("path is required")
    path = env.project._resolve_read_path(raw)  # noqa: SLF001
    if env.project._is_secret_path(path):  # noqa: SLF001
        raise PermissionError("Document path matches sensitive file policy")
    if env.project._is_protected_state_path(path):  # noqa: SLF001
        raise PermissionError("Document path targets protected internal state")
    if not path.exists() or not path.is_file():
        raise FileNotFoundError(f"Document not found: {raw}")
    return path


def _clip(text: str, max_chars: int) -> tuple[str, bool]:
    limit = max(1, min(_MAX_CHARS_LIMIT, int(max_chars or _DEFAULT_MAX_CHARS)))
    return text[:limit], len(text) > limit


def _read_pdf(path: Path) -> str | dict[str, object]:
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError:
        return _err("E_DEPENDENCY", "PDF extraction requires optional dependency: pypdf", {"path": str(path), "dependency": "pypdf"})
    reader = PdfReader(str(path))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(path: Path) -> str | dict[str, object]:
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError:
        return _err("E_DEPENDENCY", "DOCX extraction requires optional dependency: python-docx", {"path": str(path), "dependency": "python-docx"})
    document = docx.Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)


def _extract_text(args: dict[str, object], env: ToolExecutionEnv) -> dict[str, object]:
    path = _resolve_path(str(args.get("path") or ""), env)
    suffix = path.suffix.lower()
    max_chars = int(args.get("max_chars") or _DEFAULT_MAX_CHARS)
    if suffix in {".txt", ".md", ".log", ".json", ".yaml", ".yml", ".xml", ".html", ".htm"}:
        text = path.read_text(encoding="utf-8", errors="replace")
    elif suffix == ".csv":
        with path.open(newline="", encoding="utf-8", errors="replace") as handle:
            rows = list(csv.reader(handle))
        text = "\n".join(", ".join(row) for row in rows)
    elif suffix == ".pdf":
        extracted = _read_pdf(path)
        if isinstance(extracted, dict):
            return extracted
        text = extracted
    elif suffix == ".docx":
        extracted = _read_docx(path)
        if isinstance(extracted, dict):
            return extracted
        text = extracted
    else:
        return _err("E_UNSUPPORTED", f"Unsupported document type: {suffix or 'unknown'}", {"path": str(path)})
    clipped, truncated = _clip(text, max_chars)
    return _ok({"path": str(path), "format": suffix.lstrip("."), "text": clipped, "truncated": truncated, "char_count": len(text)})


def _extract_tables(args: dict[str, object], env: ToolExecutionEnv) -> dict[str, object]:
    path = _resolve_path(str(args.get("path") or ""), env)
    max_rows = max(1, min(_MAX_ROWS_LIMIT, int(args.get("max_rows") or _DEFAULT_MAX_ROWS)))
    suffix = path.suffix.lower()
    rows: list[list[str]] = []
    truncated = False
    if suffix == ".csv":
        with path.open(newline="", encoding="utf-8", errors="replace") as handle:
            for index, row in enumerate(csv.reader(handle)):
                if index >= max_rows:
                    truncated = True
                    break
                rows.append([str(cell) for cell in row])
    elif suffix == ".docx":
        try:
            import docx  # type: ignore[import-not-found]
        except ImportError:
            return _err("E_DEPENDENCY", "DOCX table extraction requires optional dependency: python-docx", {"path": str(path), "dependency": "python-docx"})
        document = docx.Document(str(path))
        for table in document.tables:
            for row in table.rows:
                if len(rows) >= max_rows:
                    truncated = True
                    return _ok({"path": str(path), "format": "docx", "rows": rows, "truncated": truncated})
                rows.append([cell.text for cell in row.cells])
    else:
        return _err("E_UNSUPPORTED", "Table extraction supports CSV and DOCX", {"path": str(path), "format": suffix.lstrip(".")})
    return _ok({"path": str(path), "format": suffix.lstrip("."), "rows": rows, "truncated": truncated})


def execute(tool_name: str, args: dict[str, object], env: ToolExecutionEnv) -> dict[str, Any]:
    if tool_name == "extract_document_text":
        return _extract_text(args, env)
    if tool_name == "extract_document_tables":
        return _extract_tables(args, env)
    raise ValueError(f"Unsupported tool: {tool_name}")
